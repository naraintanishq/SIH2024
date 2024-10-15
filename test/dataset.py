from flask import Flask, jsonify, request, send_file
from flask_pymongo import PyMongo
import requests
import pandas as pd
from io import BytesIO
from bson.json_util import dumps
from docx import Document
import xml.etree.ElementTree as ET
import bibtexparser
from scholarly import scholarly
from flask_cors import CORS

app = Flask(__name__)

# MongoDB configuration
app.config["MONGO_URI"] = "mongodb+srv://tanishqnarain:tan123@cluster0.5nwrs.mongodb.net/sihdb?retryWrites=true&w=majority&appName=cluster0"

mongo = PyMongo(app)

# Enable CORS for all routes
CORS(app, resources={r"/*": {"origins": "*"}})



# Fetch and store publications in MongoDB
@app.route('/fetch-publications', methods=['POST'])
def fetch_publications():
    data = request.json
    search_type = data.get('search_type', 'author')
    query = data.get('query')
    start_year = data.get('start_year')
    end_year = data.get('end_year')

    if not query:
        return jsonify({'error': 'Query parameter is required'}), 400

    dblp_data = fetch_dblp_data(query, search_by=search_type, start_year=start_year, end_year=end_year)
    scholar_data = search_google_scholar_and_get_bibtex(query, search_by=search_type, start_year=start_year, end_year=end_year)
    
    # Combine and store in MongoDB
    all_publications = dblp_data + scholar_data
    try:
        mongo.db.publications.delete_many({})  # Clear previous data
        mongo.db.publications.insert_many(all_publications)
    except Exception as e:
        print(f"Error inserting into MongoDB: {e}")
        return jsonify({'error': 'Failed to store publications in MongoDB'}), 500
    
    return jsonify({'message': 'Publications fetched and stored', 'count': len(all_publications)})


# DBLP API URL
DBLP_API_URL = "https://dblp.org/search/publ/api"

# Fetch DBLP data
def fetch_dblp_data(query, search_by="author", start_year=None, end_year=None):
    query = f"{search_by}:{query}"
    params = {"q": query, "h": "100"}

    response = requests.get(DBLP_API_URL, params=params)
    response.raise_for_status()
    
    publications = []
    root = ET.fromstring(response.content)
    
    for hit in root.findall(".//hit"):
        title = hit.find(".//info/title").text if hit.find(".//info/title") is not None else "N/A"
        authors = ", ".join(author.text for author in hit.findall(".//info/authors/author")) if hit.findall(".//info/authors/author") else "N/A"
        year_hit = hit.find(".//info/year").text if hit.find(".//info/year") is not None else "N/A"
        doi = hit.find(".//info/doi").text if hit.find(".//info/doi") is not None else "N/A"
        journal = hit.find(".//info/journal").text if hit.find(".//info/journal") is not None else "N/A"
        booktitle = hit.find(".//info/booktitle").text if hit.find(".//info/booktitle") is not None else "N/A"
        pages = hit.find(".//info/pages").text if hit.find(".//info/pages") is not None else "N/A"
        publisher = hit.find(".//info/publisher").text if hit.find(".//info/publisher") is not None else "N/A"
        url = hit.find(".//info/url").text if hit.find(".//info/url") is not None else "N/A"

        journal_conf = journal if journal != "N/A" else booktitle
        
        if (start_year and int(year_hit) < int(start_year)) or (end_year and int(year_hit) > int(end_year)):
            continue

        publications.append({
            "Title": title,
            "Authors": authors,
            "Year": year_hit,
            "DOI": doi,
            "Journal_Conference": journal_conf,
            "Pages": pages,
            "Publisher": publisher,
            "URL": url
        })
    
    return publications

# Fetch Google Scholar data and BibTeX
def search_google_scholar_and_get_bibtex(query, search_by="author", start_year=None, end_year=None):
    query = f"{search_by}:{query}"
    search_query = scholarly.search_pubs(query)
    
    publications = []
    
    for i, pub in enumerate(search_query):
        if i >= 5:
            break
        
        title = pub['bib'].get('title', 'N/A')
        authors = pub['bib'].get('author', 'N/A')
        year_hit = pub['bib'].get('pub_year', 'N/A')
        journal_conf = pub['bib'].get('journal', pub['bib'].get('booktitle', 'N/A'))
        pages = pub['bib'].get('pages', 'N/A')
        publisher = pub['bib'].get('publisher', 'N/A')
        bibtex_link = pub.get('url_add_scholarbib', None)
        pub_url = pub.get('url', 'N/A')
        
        if bibtex_link:
            bibtex_data = fetch_bibtex(bibtex_link)
            doi = bibtex_data.get('doi', 'N/A')
        else:
            doi = 'N/A'
        
        if (start_year and int(year_hit) < int(start_year)) or (end_year and int(year_hit) > int(end_year)):
            continue
        
        publications.append({
            "Title": title,
            "Authors": authors,
            "Year": year_hit,
            "DOI": doi,
            "Journal_Conference": journal_conf,
            "Pages": pages,
            "Publisher": publisher,
            "URL": pub_url
        })
    
    return publications

# Fetch BibTeX data
def fetch_bibtex(bibtex_url):
    try:
        response = requests.get(bibtex_url)
        response.raise_for_status()
        bibtex_content = response.text
        
        bib_database = bibtexparser.loads(bibtex_content)
        
        if bib_database.entries:
            return bib_database.entries[0]
        return {}
    except requests.RequestException as e:
        print(f"Error fetching BibTeX data: {e}")
        return {}
    except Exception as e:
        print(f"Error parsing BibTeX data: {e}")
        return {}

# Get publications from MongoDB
@app.route('/get-publications', methods=['GET'])
def get_publications():
    publications = mongo.db.publications.find()
    return jsonify(dumps(publications))

# Export publications
@app.route('/export-publications-excel', methods=['GET'])
def export_publications_excel():
    publications = mongo.db.publications.find()
    data = [pub for pub in publications]

    df = pd.DataFrame(data)
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='Publications')
    writer.save()
    output.seek(0)

    return send_file(output, attachment_filename='publications.xlsx', as_attachment=True)

@app.route('/export-publications-word', methods=['GET'])
def export_publications_word():
    publications = mongo.db.publications.find()
    data = [pub for pub in publications]

    doc = Document()
    doc.add_heading('Publications', 0)

    for pub in data:
        doc.add_paragraph(f"Title: {pub['Title']}")
        doc.add_paragraph(f"Year: {pub['Year']}")
        doc.add_paragraph(f"Type: {pub['Journal_Conference']}")
        doc.add_paragraph(f"Authors: {pub['Authors']}")
        doc.add_paragraph(f"Publisher: {pub['Publisher']}")
        doc.add_paragraph(f"DOI: {pub['DOI']}")
        doc.add_paragraph(f"URL: {pub['URL']}")
        doc.add_paragraph('')

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, attachment_filename='publications.docx', as_attachment=True)

# Query publications by year range
@app.route('/query-publications', methods=['GET'])
def query_publications():
    start_year = request.args.get('start_year')
    end_year = request.args.get('end_year')

    if not start_year or not end_year:
        return jsonify({'error': 'start_year and end_year are required'}), 400

    try:
        publications = mongo.db.publications.find({
            "Year": {
                "$gte": int(start_year),
                "$lte": int(end_year)
            }
        })
        return jsonify(dumps(publications))
    except Exception as e:
        print(f"Error querying MongoDB: {e}")
        return jsonify({'error': 'Failed to query publications'}), 500

if __name__ == '__main__':
    app.run(debug=True)
